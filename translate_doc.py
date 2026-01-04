import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from googletrans import Translator
import time
import os
import re
import copy

def merge_paragraphs(doc):
    print("Pre-processing: Merging split paragraphs (Body only)...")
    # Snapshot of paragraphs
    paragraphs = list(doc.paragraphs)
    if not paragraphs:
        return

    last_p = paragraphs[0]
    merged_count = 0
    
    terminators = ('。', '？', '！', '：', ':', ';', '；', ')', '）', '”', '"')
    
    for i in range(1, len(paragraphs)):
        curr_p = paragraphs[i]
        text = curr_p.text.strip()
        last_text = last_p.text.strip()
        
        if not text:
            last_p = curr_p
            continue
            
        if not last_text:
             last_p = curr_p
             continue
             
        # Check termination of last_p
        is_terminated = last_text.endswith(terminators)
        
        # Check next start (curr_p)
        starts_with_number = bool(re.match(r'^\s*\d+(\.|$)', text))
        
        # Check if last_p looks like a header (Short + Numbered)
        last_is_header = bool(re.match(r'^\s*\d+(\.|$)', last_text)) and len(last_text) < 40
        
        if not is_terminated and not starts_with_number and not last_is_header:
            # Merge
            sep = ""
            if last_text[-1].isascii() and text[0].isascii():
                sep = " "
                
            last_p.text = last_p.text.rstrip() + sep + text.lstrip()
            
            try:
                curr_p._element.getparent().remove(curr_p._element)
                merged_count += 1
            except Exception as e:
                print(f"Error removing paragraph: {e}")
        else:
            last_p = curr_p
            
    print(f"Merged {merged_count} paragraphs.")

def translate_paragraph_element(p, translator):
    """
    Translates a paragraph element (python-docx Paragraph object or wrapper).
    Inserts the translation as a NEW paragraph following the current one.
    """
    text = p.text.strip()
    if not text:
        return

    try:
        # Check for header case (Ends with colon)
        is_header_colon = text.endswith('\uff1a') or text.endswith(':')
        text_to_translate = text.rstrip('\uff1a:')
        
        # Translate
        # Retry logic could be added here
        result = translator.translate(text_to_translate, src='zh-cn', dest='en')
        translated_text = result.text.strip()
        
        if is_header_colon:
            # Inline translation
            p.text = f"{text_to_translate}({translated_text}):"
        else:
            # Remove numbering from translation
            cleaned_translation = re.sub(r'^[\d\.]+\s*', '', translated_text)
            
            # Insert new paragraph
            # We need to clone the style if possible, or create a new p element
            
            # Strategy: create a new paragraph element in the same parent
            # Access the low-level element
            p_elem = p._element
            parent = p_elem.getparent()
            
            # Create new p element (copy of p to keep style?)
            # copying element might be safer to keep formatting properties (indentation etc)
            # but we need to clear its content
            
            new_p_elem = copy.deepcopy(p_elem)
            
            # Clear runs in the copy
            # p children are r, pPr, etc. We want to keep pPr but remove r
            for child in list(new_p_elem):
                if child.tag.endswith('r'): # Run
                     new_p_elem.remove(child)
            
            # Now add text to the new p element
            # We can create a run element and add text
            # Or use a temporary wrappers
            
            # Simpler way with python-docx if we can: 
            # But standard doc.add_paragraph adds to body end.
            # So constructing XML manually is best for arbitrary location.
            
            # Construct a run with the text
            # We assume basic styling in runs is not critical to copy, just paragraph style
            
            # Create a run element
            run = docx.oxml.shared.OxmlElement('w:r')
            t = docx.oxml.shared.OxmlElement('w:t')
            t.text = cleaned_translation
            run.append(t)
            new_p_elem.append(run)
            
            # Insert after current p
            parent.insert(parent.index(p_elem) + 1, new_p_elem)
            
    except Exception as e:
        print(f"  Error translating '{text[:20]}...': {e}")

def translate_and_format(input_path, output_path):
    print(f"Loading {input_path}...")
    doc = docx.Document(input_path)
    
    # 1. Merge Paragraphs (Body)
    merge_paragraphs(doc)
    
    translator = Translator()
    
    # helper for processing lists of paragraphs
    def process_paragraph_list(paragraph_list):
        # iterate copy
        for p in list(paragraph_list):
            if p.text.strip():
                translate_paragraph_element(p, translator)
                time.sleep(0.2)

    # 2. Translate Body
    print("Translating Body...")
    process_paragraph_list(doc.paragraphs)
    
    # 3. Translate Headers/Footers
    print("Translating Headers/Footers...")
    for section in doc.sections:
        process_paragraph_list(section.header.paragraphs)
        process_paragraph_list(section.footer.paragraphs)
        
    # 4. Translate Tables
    print("Translating Tables...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraph_list(cell.paragraphs)
                
    # 5. Translate Text Boxes
    print("Translating Text Boxes...")
    # Iterate all txbxContent
    body = doc.element.body
    txbx_contents = body.findall('.//' + qn('w:txbxContent')) # This might need full traversal if nested
    # findall only looks at direct children if not using .// but .// works on ElementTree
    # lxml/ElementTree support xpath-like
    
    # If iterating over all elements:
    count_txbx = 0
    for txbx in body.iter(qn('w:txbxContent')):
        count_txbx += 1
        # Find paragraphs inside
        ps = list(txbx.iter(qn('w:p')))
        for p_elem in ps:
            # Wrap in Paragraph
            # We assume 'parent' arg as None is okay for text access/manipulation if we don't access complex styles
            # But strictly, Paragraph expects a parent.
            # We can use the txbx as parent proxy if needed, or just None.
            # Testing shows Paragraph(elem, None) usually works for .text and basic ops
            try:
                p_obj = Paragraph(p_elem, None) 
                if p_obj.text.strip():
                     translate_paragraph_element(p_obj, translator)
                     time.sleep(0.2)
            except Exception as e:
                print(f"  Error in text box p: {e}")
                
    print(f"Processed {count_txbx} text boxes.")

    print(f"Saving to {output_path}...")
    doc.save(output_path)
    print("Done.")

if __name__ == "__main__":
    input_file = "document_cn.docx"
    output_file = "translated_document.docx"
    
    if os.path.exists(input_file):
        translate_and_format(input_file, output_file)
    else:
        print("Input file not found.")
